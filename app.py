# app.py
# pip install streamlit python-docx streamlit-drawable-canvas pillow

import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime, timedelta
import random
import tempfile
from PIL import Image
from streamlit_drawable_canvas import st_canvas

# ---------------- CONFIG ----------------
st.set_page_config(page_title="Employee Contract Generator", layout="centered")

TEMPLATE_FILE = "Contract_template.docx"
POSITION = "Cleaning Services Employee"
PAY_RATE = "$30.35 per hour"
COMPANIES = {
    "Cockburn Gateways": {
        "company_name": "Finolex Auto Pty Ltd",
        "company_address": "Cockburn Gateways",
        "manager_name": "Harpreet Singh",
        "manager_phone": "0414409284",
    },
    "Galleria": {
        "company_name": "Ks & K Pty Ltd",
        "company_address": "Star Car Wash - Galleria, Galleria Shopping Centre, 4 Collier Rd, Morley WA 6062",
        "manager_name": "Kuldeep Kuldeep",
        "manager_phone": "0497580120",
    },
    "Innaloo": {
        "company_name": "Star Car Wash",
        "company_address": "Westfield Innaloo (Level 1 car park near Kmart Service Centre), Cnr Scarborough Beach Rd & Ellen Stirling Blvd, Innaloo WA 6021",
        "manager_name": "Raj Partap Singh Dhanju",
        "manager_phone": "0470382410",
    },
}

# ---------------- HELPERS ----------------
def format_date(dt):
    return dt.strftime("%d/%m/%Y")

def random_contract_date(start_dt):
    # random 2 to 7 days before start date
    days_before = random.randint(2, 7)
    return start_dt - timedelta(days=days_before)

def replace_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)

    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True

    if changed:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full_text

def replace_everywhere(doc, replacements):
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)

def insert_signature(doc, image_path):
    for p in doc.paragraphs:
        if "{{signature}}" in p.text:
            p.text = p.text.replace("{{signature}}", "")
            run = p.add_run()
            run.add_picture(image_path, width=Inches(2.2))
            return

def generate_contract(data, company, signature_path=None):
    doc = Document(TEMPLATE_FILE)

    replacements = {
        "{{position}}": data["position"],
        "{{employee_name}}": data["employee_name"],
        "{{employee_address}}": data["employee_address"],
        "{{start_date}}": data["start_date"],
        "{{date_today}}": data["date_today"],
        "{{company_name}}": company["company_name"],
        "{{company_address}}": company["company_address"],
        "{{manager_name}}": company["manager_name"],
        "{{manager_phone}}": company["manager_phone"],
    }

    replace_everywhere(doc, replacements)

    if signature_path:
        insert_signature(doc, signature_path)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ---------------- UI ----------------
st.title("Employee Contract Generator")

st.markdown(f"**Position:** {POSITION}")
st.markdown(f"**Pay Rate:** {PAY_RATE}")
company_key = st.selectbox("Select Company", list(COMPANIES.keys()))
company = COMPANIES[company_key]

with st.form("contract_form"):
    employee_name = st.text_input("Full Name")
    employee_address = st.text_area("Address")
    start_date = st.date_input("Employment Start Date")

    st.markdown("### Draw Signature")
    canvas = st_canvas(
        fill_color="rgba(255,255,255,0)",
        stroke_width=2,
        stroke_color="black",
        background_color="white",
        height=160,
        width=500,
        drawing_mode="freedraw",
        key="signature_box",
    )

    submitted = st.form_submit_button("Generate Contract")

# ---------------- PROCESS ----------------
if submitted:
    if not employee_name.strip() or not employee_address.strip():
        st.error("Please complete all required fields.")
        st.stop()

    # Start date
    start_dt = datetime.combine(start_date, datetime.min.time())

    # Random contract date 2-7 days before start
    contract_dt = random_contract_date(start_dt)

    # Signature save
    sig_path = None
    if canvas.image_data is not None:
        img = Image.fromarray(canvas.image_data.astype("uint8"))
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        img.save(tmp.name)
        sig_path = tmp.name

    data = {
        "position": POSITION,
        "employee_name": employee_name.strip(),
        "employee_address": employee_address.strip(),
        "start_date": format_date(start_dt),
        "date_today": format_date(contract_dt),
    }

    output = generate_contract(data, company, sig_path)

    filename = employee_name.strip().replace(" ", "_") + "_Contract.docx"

    st.success("Contract generated successfully.")
    st.download_button(
        "Download Contract",
        data=output,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )